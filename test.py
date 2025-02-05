from fastapi import FastAPI, File, UploadFile, Form
from fastapi.responses import HTMLResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi import Request
import os
import ast
import re
from unstructured.partition.docx import partition_docx

from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
import os
from docx import Document
doc = Document()
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from openai import AzureOpenAI

def config():
    load_dotenv()

endpoint = "https://enrollment-copilot-formrecognizer-dev.cognitiveservices.azure.com/"  
config()
document_analysis_client = DocumentAnalysisClient(endpoint=endpoint, credential=AzureKeyCredential(os.getenv('api_key')))

# Initialize FastAPI app and Jinja2 templates
app = FastAPI()

templates = Jinja2Templates(directory="templates")

# Serve static files like CSS, JS, etc. (if needed for styling)
# app.mount("/static", StaticFiles(directory="static"), name="static")

# Endpoint to show the HTML form
@app.get("/", response_class=HTMLResponse)
async def upload_form(request: Request):
    return templates.TemplateResponse("first.html", {"request": request})

# Endpoint to handle the file upload
@app.post("/upload/")
async def upload_file(request: Request,file: UploadFile = File(...),file1: UploadFile = File(...)):
    # Save the uploaded file to the server
   
    file_location = f"uploads/{file.filename}"
    # print(await file.read())
    with open(file_location, "wb") as f:
        f.write(await file.read())
    with open(file_location, "rb") as f:
    # print('f',f)
        poller = document_analysis_client.begin_analyze_document("prebuilt-read", document=f, locale="en-US")
    # Get the result of the analysis
        result = poller.result()
        file2 = open("guide.txt", "w")
        
        for page in result.pages:
            txt_guide = ' '
            print(f"Page number: {page.page_number}")
            for word in page.words:
                txt_guide = txt_guide + ' ' + word.content
            file2.write(txt_guide) 
            file2.write("\n")
        file2.close()
    file2 = open(r"guide.txt", "r+")


    guide =file2.read()
    file2.close()
    print('txt_guide')
    print(guide)


    file_location = f"uploads/{file1.filename}"
    with open(file_location, "wb") as f:
        f.write(await file1.read())
    general = partition_docx(file_location)
    txt_doc = ' '
    for gen in general:
        # print(gen.text)
        txt_doc = txt_doc + '\n' + gen.text
    print('txt_doc')
    print(txt_doc)

    
    openai = AzureOpenAI(
    api_version='2024-02-01',azure_endpoint="https://openai-wu-genai-mmp-d.openai.azure.com/",azure_deployment='gpt-4o',api_key="f92989b781a64de6abaae07cf062b02e")
    
    chat_completion1 = openai.chat.completions.create(
                
                model='gpt-4o',
                messages=[
        {
        "role": "system",
        "content": """
        Your task is to go through the provided guide and extract the guidelines that relates to formatting the bibliography 
        """
        },
        {
        "role": "user",
        "content": guide
        }
                ], temperature=0
    )

    bib_guide = chat_completion1.choices[0].message.content
    print('bib',bib_guide)
    
    chat_completion = openai.chat.completions.create(
                
                model='gpt-4o',
                messages=[
        {
        "role": "system",
        "content": """Your task is to extract Course name and all the 
        bibliography present in the content. 
        Bibliography should includes all the media details along with the reading materials
        Strictly provide only the bibliography details only.Don't group them.
        Arrange the bibliography in ascending order
        Give the output in a json format : {'course':'course name','bibliography':[List of bibliographies]}
        """
        },
        {
        "role": "user",
        "content": txt_doc
        }
                ], temperature=0
    )

    task = chat_completion.choices[0].message.content
    start = re.search('{',task).start()
    end = re.search('}',task).start()
    task = task[start:end+1]
    task1 = ast.literal_eval(task)
    a = task1['bibliography']
    task2 = '\n'.join(a)
    print('task2')
    print(task2)
    course = task1['course']
    print('first openai call done')
    # print(txt_doc)
    messages = [
        {
        "role": "assistant",
        #   "content": f"""Your task is to go through the guidelines given below and format the bibliographies 
        "content": f"""Please format the provided bibliographies using the below rules/guidelines.
        You have to abide by all the rules
            
        Rules related to Bibliography in the Walden PDCD Editorial Style Handbook:

        {bib_guide}

        Do not remove any bibliography. Do not replicate the bibliography
        
        Arrange the bibliographies in ascending order
        Only extract the bibliography, nothing else
        
        """
        },
        {
        "role": "user",
        "content":  task2
        }
    ]

    openai = AzureOpenAI(
    api_version='2024-02-01',azure_endpoint="https://openai-wu-genai-mmp-d.openai.azure.com/",azure_deployment='o1-mini',api_key="f92989b781a64de6abaae07cf062b02e")

    chat_completion = openai.chat.completions.create(
                
                model='o1-mini',
                messages= messages,
    
    )
    print('second opi done')
    print(chat_completion.choices[0].message.content)
    title = doc.add_heading('BIBLIOGRAPHY', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_heading(course, level=2)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph = doc.add_paragraph(chat_completion.choices[0].message.content)
    run = paragraph.runs[0]
    run.bold = True
    print('doc created')
    doc.save('uploads/test.docx')

    general = partition_docx('uploads/test.docx')
    txt_final = ' '
    for gen in general:
        # print(gen.text)
        txt_final = txt_final + '\n' + gen.text
    # print('final_doc')
    # print(txt_final)



    return templates.TemplateResponse("first.html", {"request": request, "filename": file1.filename, "text": txt_final})


# Create the uploads folder if it doesn't exist
if not os.path.exists("uploads"):
    os.makedirs("uploads")
